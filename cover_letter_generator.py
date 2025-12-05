#!/usr/bin/env python3
"""
Cover Letter Generator GUI (Tkinter) with ChatGPT integration (OpenAI Python client).

- Uses the OpenAI Python client (from openai import OpenAI).
- Set OPENAI_API_KEY in environment before running.
- Provide: resume (.txt or .docx), base cover letter (.txt or .docx), job posting text (paste),
  company, position, and applicant name. Click "Generate Preview" to have ChatGPT produce
  a complete, polished cover letter (no bullets, no lists) using the style/tone of your base letter.

Requirements:
- python-docx (optional, required only if you want to read/write .docx)
- openai (the OpenAI Python client available via `pip install openai`)
"""

import os
import re
import collections
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from tkinter.scrolledtext import ScrolledText

# ChatGPT client (Option A)
from openai import OpenAI

# try to import python-docx (optional)
try:
    import docx
except Exception:
    docx = None

# initialize OpenAI client using environment variable
OPENAI_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_KEY:
    # We won't crash at import, but will check before calling the API.
    client = None
else:
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))



def read_text_file(path):
    """Read .txt or .docx content (docx optional). Return a string."""
    if not path:
        return ''
    p = str(path)
    if p.lower().endswith('.docx'):
        if docx is None:
            return ''
        try:
            d = docx.Document(p)
            paras = [para.text for para in d.paragraphs if para.text.strip()]
            return '\n\n'.join(paras)
        except Exception:
            return ''
    else:
        try:
            with open(p, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception:
            return ''


def save_as_docx(text, path):
    """Save plain text into a simple .docx file (paragraphs separated by blank lines)."""
    if docx is None:
        return False, 'python-docx not installed'
    try:
        d = docx.Document()
        for para in text.split('\n\n'):
            d.add_paragraph(para.strip())
        d.save(path)
        return True, None
    except Exception as e:
        return False, str(e)


def save_as_txt(content, path):
    try:
        with open(path, 'w', encoding='utf-8') as f:
            f.write(content)
        return True, None
    except Exception as e:
        return False, str(e)


def chatgpt_generate_cover_letter(base_letter, resume_text, job_posting, company, position, name):
    """
    Uses the OpenAI Chat Completions API (client.chat.completions.create) to generate
    a complete cover letter.

    Option 1 behavior:
    - ChatGPT rewrites the entire cover letter from scratch using the tone/style/structure
      of the provided base letter. Produces full paragraphs (no bullets/lists).
    """

    if client is None:
        raise RuntimeError("OpenAI client not initialized. Please set the OPENAI_API_KEY environment variable.")

    # System-level instructions describing the assistant role
    system_prompt = (
        "You are an expert career-writing assistant. Your job is to produce polished, professional, and "
        "persuasive cover letters. Use the tone, style, and structure of the provided base letter as a guide, "
        "but rewrite the letter entirely from scratch. Produce 3–5 well-formed paragraphs, including a clear opening "
        "that mentions the company and role, tailored middle paragraphs that reference the candidate's resume and the job posting, "
        "and a strong closing paragraph. Do NOT produce bullet points, numbered lists, or sections. Keep the letter concise and natural."
    )

    # Compose the user prompt with the base letter, resume, posting, and metadata.
    # Keep the prompt reasonably sized (truncate resume or posting if extremely large).
    def _truncate_for_prompt(text, max_chars=3500):
        if not text:
            return ''
        if len(text) <= max_chars:
            return text
        # keep start and end for context
        head = text[:max_chars//2]
        tail = text[-(max_chars//2):]
        return head + '\n\n...TRUNCATED...\n\n' + tail

    base_letter_snip = _truncate_for_prompt(base_letter, max_chars=2000)
    resume_snip = _truncate_for_prompt(resume_text, max_chars=3000)
    posting_snip = _truncate_for_prompt(job_posting, max_chars=3500)

    user_prompt = f"""
Write a complete professional cover letter using the base letter as a style guide.

Rules:
- Rewrite the letter entirely using the tone and structure of the base letter.
- Mention Company ("{company}") and Position ("{position}") in the opening paragraph.
- Use the resume details to represent the candidate accurately, but do not invent specific claims (treat resume text as authoritative).
- Use the job posting to tailor which skills/experiences to emphasize.
- Produce 3-5 full paragraphs total.
- No bullet points, no lists, no placeholders like {{tailored_points}}.
- Keep the letter suitable for a professional job application (formal, confident, concise).
- If the resume or base letter provides the applicant name, sign with the provided name: "{name}".

Base letter (style guide):
{base_letter_snip}

Resume:
{resume_snip}

Job posting:
{posting_snip}

Generate only the final cover letter text (do not include commentary, JSON, or extra metadata).
"""

    # Call the Chat Completions endpoint
    # Use a capable model; change model name if you prefer
    model_name = "gpt-4.1"

    try:
        response = client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.6,
            max_tokens=900,
        )
    except Exception as e:
        # surface API errors to the caller
        raise RuntimeError(f"OpenAI API error: {e}")

    # Extract content safely
    text = None
    try:
        # new OpenAI python client response structure:
        text = response.choices[0].message["content"]
    except Exception:
        try:
            text = response.choices[0].message.content
        except Exception:
            # fallback: try 'text' or 'delta' patterns
            try:
                text = response.choices[0].text
            except Exception:
                text = None

    if not text:
        raise RuntimeError("Failed to retrieve content from OpenAI response.")

    return text.strip()


class CoverLetterApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title('Cover Letter Generator (ChatGPT)')
        self.geometry('900x700')
        self._build_ui()

    def _build_ui(self):
        frm = ttk.Frame(self, padding=10)
        frm.pack(fill='both', expand=True)

        # Row 0: Resume and base letter selection
        ttk.Label(frm, text='Resume file:').grid(row=0, column=0, sticky='w')
        self.resume_var = tk.StringVar()
        ttk.Entry(frm, textvariable=self.resume_var, width=60).grid(row=0, column=1, sticky='w')
        ttk.Button(frm, text='Browse', command=self.browse_resume).grid(row=0, column=2, sticky='w')

        ttk.Label(frm, text='Base cover letter:').grid(row=1, column=0, sticky='w')
        self.base_var = tk.StringVar()
        ttk.Entry(frm, textvariable=self.base_var, width=60).grid(row=1, column=1, sticky='w')
        ttk.Button(frm, text='Browse', command=self.browse_base).grid(row=1, column=2, sticky='w')

        # Row 2: Company, position, name
        ttk.Label(frm, text='Company:').grid(row=2, column=0, sticky='w')
        self.company_var = tk.StringVar()
        ttk.Entry(frm, textvariable=self.company_var, width=30).grid(row=2, column=1, sticky='w')

        ttk.Label(frm, text='Position:').grid(row=2, column=2, sticky='w')
        self.position_var = tk.StringVar()
        ttk.Entry(frm, textvariable=self.position_var, width=30).grid(row=2, column=3, sticky='w')

        ttk.Label(frm, text='Your name:').grid(row=2, column=4, sticky='w')
        self.name_var = tk.StringVar()
        ttk.Entry(frm, textvariable=self.name_var, width=25).grid(row=2, column=5, sticky='w')

        # Row 3: Job posting textarea
        ttk.Label(frm, text='Job posting (paste text here):').grid(row=3, column=0, columnspan=6, sticky='w', pady=(8, 0))
        self.job_text = ScrolledText(frm, width=110, height=12)
        self.job_text.grid(row=4, column=0, columnspan=6, pady=(0, 8))

        # Row 5: Output filename and folder
        ttk.Label(frm, text='Output filename (no extension):').grid(row=5, column=0, sticky='w')
        self.out_filename_var = tk.StringVar(value='Cover_Letter')
        ttk.Entry(frm, textvariable=self.out_filename_var, width=40).grid(row=5, column=1, sticky='w')

        ttk.Label(frm, text='Output folder:').grid(row=5, column=2, sticky='w')
        self.out_folder_var = tk.StringVar(value=os.getcwd())
        ttk.Entry(frm, textvariable=self.out_folder_var, width=40).grid(row=5, column=3, sticky='w')
        ttk.Button(frm, text='Browse', command=self.browse_out_folder).grid(row=5, column=4, sticky='w')

        # Row 6: Buttons
        gen_btn = ttk.Button(frm, text='Generate Preview', command=self.generate_preview)
        gen_btn.grid(row=6, column=1, sticky='w', pady=10)
        save_btn = ttk.Button(frm, text='Save Cover Letter', command=self.save_cover_letter)
        save_btn.grid(row=6, column=2, sticky='w', pady=10)
        copy_btn = ttk.Button(frm, text='Copy Preview to Clipboard', command=self.copy_preview)
        copy_btn.grid(row=6, column=3, sticky='w', pady=10)

        # Row 7: Preview area
        ttk.Label(frm, text='Preview:').grid(row=7, column=0, columnspan=6, sticky='w')
        self.preview = ScrolledText(frm, width=110, height=18)
        self.preview.grid(row=8, column=0, columnspan=6, pady=(0, 8))

        # Help note
        help_text = (
            "Placeholders supported in base cover letter (style guide only): {company}, {position}, {name}.\n"
            "This app will ask ChatGPT to rewrite a complete cover letter using the base letter's tone and your resume + job posting.\n"
            "Supported file types: .txt and .docx for resume and base letter."
        )
        ttk.Label(frm, text=help_text, foreground='gray').grid(row=9, column=0, columnspan=6, sticky='w')

    def browse_resume(self):
        p = filedialog.askopenfilename(filetypes=[('Text or DOCX', '*.txt;*.docx'), ('All files', '*.*')])
        if p:
            self.resume_var.set(p)

    def browse_base(self):
        p = filedialog.askopenfilename(filetypes=[('Text or DOCX', '*.txt;*.docx'), ('All files', '*.*')])
        if p:
            self.base_var.set(p)

    def browse_out_folder(self):
        p = filedialog.askdirectory()
        if p:
            self.out_folder_var.set(p)

    def generate_preview(self):
        resume_path = self.resume_var.get()
        base_path = self.base_var.get()
        job_posting = self.job_text.get('1.0', 'end').strip()
        company = self.company_var.get().strip()
        position = self.position_var.get().strip()
        name = self.name_var.get().strip()

        resume_text = read_text_file(resume_path)
        if not resume_text and resume_path:
            messagebox.showwarning('Resume', 'Unable to read resume file (unsupported format or empty). Proceeding without resume content.')

        base_letter = read_text_file(base_path)
        if not base_letter and base_path:
            messagebox.showwarning('Base letter', 'Unable to read base cover letter file (unsupported format or empty). Using a simple default.')
            base_letter = (
                "Dear Hiring Manager,\n\n"
                "I am writing to express my interest in the {position} role at {company}. "
                "I believe my background and experience make me a strong fit for this opportunity.\n\n"
                "Sincerely,\n{name}"
            )

        if not job_posting:
            messagebox.showwarning('Job Posting', 'Please paste the job posting in the box before generating.')
            return

        # Call ChatGPT to produce the full cover letter
        try:
            final_letter = chatgpt_generate_cover_letter(
                base_letter=base_letter,
                resume_text=resume_text,
                job_posting=job_posting,
                company=company,
                position=position,
                name=name
            )
        except Exception as e:
            messagebox.showerror("ChatGPT Error", f"Failed to generate cover letter:\n{e}")
            return

        # Insert into preview
        self.preview.delete('1.0', 'end')
        self.preview.insert('1.0', final_letter)

    def copy_preview(self):
        text = self.preview.get('1.0', 'end').strip()
        if not text:
            messagebox.showinfo('Copy', 'Nothing in preview to copy.')
            return
        self.clipboard_clear()
        self.clipboard_append(text)
        messagebox.showinfo('Copy', 'Preview copied to clipboard.')

    def save_cover_letter(self):
        text = self.preview.get('1.0', 'end').strip()
        if not text:
            messagebox.showwarning('Save', 'No generated cover letter to save. Generate a preview first.')
            return

        out_folder = self.out_folder_var.get() or os.getcwd()
        os.makedirs(out_folder, exist_ok=True)
        filename = self.out_filename_var.get().strip() or 'Cover_Letter'

        choice = messagebox.askquestion('Format', 'Save as Word (.docx)? Click Yes for .docx, No for .txt')
        if choice == 'yes':
            out_path = os.path.join(out_folder, filename + '.docx')
            ok, err = save_as_docx(text, out_path)
        else:
            out_path = os.path.join(out_folder, filename + '.txt')
            ok, err = save_as_txt(text, out_path)

        if ok:
            messagebox.showinfo('Saved', f'Saved to {out_path}')
            try:
                os.startfile(out_path)
            except Exception:
                pass
        else:
            messagebox.showerror('Error', f'Failed to save: {err}')


if __name__ == '__main__':
    # On startup, remind user if API key is missing
    if not OPENAI_KEY:
        root = tk.Tk()
        root.withdraw()
        messagebox.showwarning(
            "OPENAI_API_KEY not set",
            "OPENAI_API_KEY environment variable is not set. Please set it before using the ChatGPT functionality.\n\n"
            "See script comments for instructions."
        )
        root.destroy()

    app = CoverLetterApp()
    app.mainloop()
